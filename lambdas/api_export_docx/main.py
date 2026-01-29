import base64
from io import BytesIO
import json
import logging
import os

import boto3
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt

logger = logging.getLogger()
logger.setLevel(logging.INFO)

dynamodb = boto3.resource("dynamodb")
TABLE_NAME = os.environ.get("DYNAMODB_TABLE_NAME")
if not TABLE_NAME:
    raise RuntimeError("Missing env var DYNAMODB_TABLE_NAME")
table = dynamodb.Table(TABLE_NAME)


def _response(status_code: int, body: dict, *, headers: dict | None = None):
    return {
        "statusCode": status_code,
        "headers": {"Content-Type": "application/json", **(headers or {})},
        "body": json.dumps(body),
    }


def _set_body_font(p, *, bold=False, italic=False):
    try:
        run = p.runs[0]
    except IndexError:
        run = p.add_run("")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bool(bold)
    run.italic = bool(italic)


def _append_verdict_form(document: Document, verdict_form: dict | None) -> None:  # noqa: PLR0912, PLR0915
    if not isinstance(verdict_form, dict):
        return
    title_p = document.add_paragraph("JURY VERDICT FORM")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_body_font(title_p, bold=True)

    # Case caption with actual party names (centered)
    def _fmt_party(n):
        if isinstance(n, list):
            return ", ".join([str(x).strip() for x in n if str(x).strip()])
        return str(n or "").strip()

    p_name = _fmt_party(verdict_form.get("plaintiff"))
    d_name = _fmt_party(verdict_form.get("defendant"))
    caption_text = None
    if p_name and d_name:
        caption_text = f"{p_name} v. {d_name}"
    elif p_name:
        caption_text = p_name
    elif d_name:
        caption_text = d_name

    if caption_text:
        hdr = document.add_paragraph(caption_text)
        hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_body_font(hdr, bold=True)

    sections = verdict_form.get("sections") or []
    for sec in sections:
        # Section title
        st = str((sec or {}).get("title") or "").strip()
        if st:
            sp = document.add_paragraph(st.upper())
            sp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_body_font(sp, bold=True)

        if (sec or {}).get("type") == "apportionment":
            instr = (sec or {}).get("instruction") or "If you found negligence, state percentage of fault. Total must equal 100%."  # noqa: E501
            ip = document.add_paragraph(instr)
            _set_body_font(ip)
            parties = (sec or {}).get("parties") or []
            for name in parties:
                line = document.add_paragraph(f"{name}: ______ %")
                _set_body_font(line)
            continue

        # Standard claim blocks
        for claim in (sec or {}).get("claims") or []:
            ct = str((claim or {}).get("claim_title") or "").strip()
            if ct:
                cp = document.add_paragraph(ct)
                _set_body_font(cp, bold=True)

            # Questions
            for q in (claim or {}).get("questions") or []:
                qnum = q.get("number")
                qtext = q.get("text") or ""
                qtype = q.get("type") or ""
                # Ensure question numbers render even if stored as Decimal or string
                try:
                    qnum_int = int(qnum) if qnum is not None else None
                except Exception:
                    qnum_int = None
                if qnum_int is not None:
                    qp = document.add_paragraph(f"{qnum_int}. {qtext}")
                else:
                    qp = document.add_paragraph(str(qtext))
                _set_body_font(qp)

                rt = (q.get("response_type") or "").lower()
                if rt == "yes_no":
                    yn = document.add_paragraph("Yes _____    No _____")
                    _set_body_font(yn)
                if qtype == "setoff" and (q.get("followup") or ""):
                    fu = document.add_paragraph(q.get("followup"))
                    _set_body_font(fu)

                # Instruction-only row (no number)
                if qtype == "instruction" and not q.get("number"):
                    ip = document.add_paragraph(q.get("text") or "")
                    _set_body_font(ip)

            # Damages block (if present)
            dmg = (claim or {}).get("damages") or None
            if isinstance(dmg, dict):
                dnum = dmg.get("number")
                try:
                    dnum_int = int(dnum) if dnum is not None else None
                except Exception:
                    dnum_int = None
                dtext = dmg.get("question") or "Damages:"
                if dnum_int is not None:
                    p = document.add_paragraph(f"{dnum_int}. {dtext}")
                else:
                    p = document.add_paragraph(dtext)
                _set_body_font(p)
                amt = document.add_paragraph("$________________________")
                _set_body_font(amt)
                pun = dmg.get("punitive") or None
                if isinstance(pun, dict):
                    pnum = pun.get("number")
                    try:
                        pnum_int = int(pnum) if pnum is not None else None
                    except Exception:
                        pnum_int = None
                    ptext = pun.get("question") or "Punitive damages:"
                    if pnum_int is not None:
                        pp = document.add_paragraph(f"{pnum_int}. {ptext}")
                    else:
                        pp = document.add_paragraph(ptext)
                    _set_body_font(pp)
                    yn = document.add_paragraph("Yes _____    No _____")
                    _set_body_font(yn)
                    if pun.get("followup"):
                        fu = document.add_paragraph(pun["followup"])
                        _set_body_font(fu)

    # Foreperson signature block at the end of the verdict form
    spacer = document.add_paragraph("")
    _set_body_font(spacer)
    spacer2 = document.add_paragraph("")
    _set_body_font(spacer2)

    sig = document.add_paragraph("Foreperson Signature: ______________________________")
    _set_body_font(sig)
    date = document.add_paragraph("Date: ____________________")
    _set_body_font(date)


def build_docx(instructions: list[dict], party_type: str, *, verdict_form: dict | None) -> bytes:
    document = Document()

    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for i, instruction in enumerate(instructions):
        jury_instruction_item_number_paragraph = document.add_paragraph(
            f'{party_type}’S REQUESTED JURY INSTRUCTION NO. {i + 1}\n'  # noqa: RUF001
        )

        run = jury_instruction_item_number_paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True

        jury_instruction_title_paragraph = document.add_paragraph(instruction['title'].upper())
        jury_instruction_title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = jury_instruction_title_paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True

        jury_instruction_content_paragraph = document.add_paragraph(instruction['customized_text'])

        run = jury_instruction_content_paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        if instruction['number'].startswith('CUSTOM-DEFAMATION-'):
            instruction_number = f'Custom Jury Instruction {instruction["number"]}'
        else:
            instruction_number = f'Florida Standard Jury Instruction {instruction["number"]}'

        jury_instruction_number_paragraph = document.add_paragraph(instruction_number)

        run = jury_instruction_number_paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.italic = True

        jury_instruction_status_paragraph = document.add_paragraph('''
Granted ___________
Denied ___________
Withdrawn ___________'''.strip())

        run = jury_instruction_status_paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        run.add_break(WD_BREAK.PAGE)

    # Append verdict form after instructions if present
    if verdict_form:
        _append_verdict_form(document, verdict_form)

    out = BytesIO()
    document.save(out)
    return out.getvalue()


def lambda_handler(event, context):
    try:
        path_params = event.get("pathParameters") or {}
        job_id = path_params.get("id") or path_params.get("job_id")
        if not job_id:
            return _response(400, {"error": "Missing id in path"})

        res = table.get_item(Key={"jury_instruction_id": job_id})
        item = res.get("Item")
        if not item:
            return _response(404, {"error": "Record not found"})

        status = item.get("status")
        if status != "COMPLETE":
            return _response(409, {"error": f"Record is not complete (status={status})"})

        # Log party_type from saved config for downstream usage/inspection
        cfg = item.get("config") or {}
        party_type = str(cfg.get("party_type", "")).upper() if isinstance(cfg, dict) else ""

        instructions = item.get("jury_instructions_text") or []
        if not isinstance(instructions, list):
            return _response(500, {"error": "Invalid instructions format"})

        verdict_form = item.get("verdict_form")
        docx_bytes = build_docx(instructions, party_type, verdict_form=verdict_form)
        b64 = base64.b64encode(docx_bytes).decode("ascii")

        filename = f"JuryInstructions-{job_id}.docx"
        return {
            "statusCode": 200,
            "headers": {
                "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "Content-Disposition": f"attachment; filename=\"{filename}\"",
            },
            "isBase64Encoded": True,
            "body": b64,
        }
    except Exception as e:
        logger.exception("Failed to generate docx")
        return _response(500, {"error": f"Failed to generate document: {e}"})
